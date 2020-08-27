import dash
import dash_core_components as dcc
import dash_html_components as html
from dash.dependencies import Input, Output
from dash import no_update
import docx
import ReferencesAutoFormat as RAF

external_stylesheets = ['https://codepen.io/chriddyp/pen/bWLwgP.css']

app = dash.Dash(__name__, external_stylesheets=external_stylesheets)

''' ### DASH GUI MODULE ###
Interactive input with example reference
Example: 
Reference: [[Fedorov, D. G.]; [Gordon, M. S.] ] 
[A study of the relative importance of one and two-electron contributions to spin−orbit coupling.]
[J. Chem. Phys.] [2000], [112], [5611−5623].

'''

app.layout = html.Div([

    html.H1('ReFFormator'),

    html.H4('Order editor'),
    html.Div([
        html.Label('Names Order'),
        dcc.Dropdown(
            id='nameorder',
            options=[
                {'label': 'Name', 'value': 'NAME'},
                {'label': 'Second Name', 'value': 'SECNAME'},
                {'label': 'Surname', 'value': 'SURNAME'}
            ],
            value=['NAME', 'SECNAME', 'SURNAME'],
            multi=True,
            style={'width': '85%'}
        ),
        #dcc.Markdown(id='nameorderout'),

        html.Label('New reference style'),
        dcc.Dropdown(
            id='referenceorder',
            options=[
                {'label': 'Names', 'value': 'NAMES'},
                {'label': 'Title', 'value': 'TITLE'},
                {'label': 'Journal', 'value': 'JOURNAL'},
                {'label': 'Year', 'value': 'YEAR'},
                {'label': 'Volume', 'value': 'VOLUME'},
                {'label': 'Issue', 'value': 'ISSUE'},
                {'label': 'Page Range', 'value': 'PAGERANGE'},
                {'label': 'DOI', 'value': 'DOI'}
            ],
            value=['NAMES', 'TITLE', 'JOURNAL', 'YEAR', 'VOLUME', 'ISSUE', 'PAGERANGE', 'DOI'],
            multi=True,
            style={'width': '85%'}
        ),
        #dcc.Markdown(id='referenceorderout')
    ]),

    # delimiters
    html.H4('Delimiters'),
    html.Div([

        html.Div([
            html.Label('Names Delimiter'),
            dcc.Input(id='namesdelimiter', value=' # #, ', type='text',
                    style={'width': '100%'}),
        ]),
        html.Div([
            html.Label('Reference Delimiters'),
            dcc.Input(id='itemsdelimiter', value='; #; #; #; #; #; #; #;', type='text',
                    style={'width': '100%'}),
        ]),
        html.Label('Use # to split the delimeters'),
    ],
    style={'display': 'inline-block', 'columns': '3'}),

    # Names
    html.H4('Names format'),
    html.Div([
        html.Div([
            html.Label('Name'),
            dcc.Input(id='name', value='John', type='text',
                        style={'width': '100%'}),
            dcc.RadioItems(
                    id='nametype',
                    options=[
                        {'label': 'Full Name', 'value': 'full'},
                        {'label': 'Short Name', 'value': 'short'}
                    ],
                    value='full'
                ),
            #dcc.Markdown(id='nameout')
        ]),

        html.Div([
            html.Label('Second Name'),
            dcc.Input(id='secname', value='Johnovich', type='text',
                        style={'width': '100%'}),
            dcc.RadioItems(
                    id='secnametype',
                    options=[
                        {'label': 'Full Name', 'value': 'full'},
                        {'label': 'Short Name', 'value': 'short'}
                    ],
                    value='full'
                ),
            #dcc.Markdown(id='secnameout')
        ]),

        html.Div([
            html.Label('Surname'),
            dcc.Input(id='surname', value='Smith', type='text',
                        style={'width': '100%'}),
            dcc.RadioItems(
                    id='surnametype',
                    options=[
                        {'label': 'Full Name', 'value': 'full'},
                        {'label': 'Short Name', 'value': 'short'}
                    ],
                    value='full'
                ),
            #dcc.Markdown(id='surnameout')
        ]),
    ],
    style={'display': 'inline-block', 'columns': '3'}),

    html.H4('Reference settings'),
    html.Label('Title'),
    html.Div([
        dcc.Input(id='Title', value='Article title', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='titlestyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[],
            ),
            #dcc.Markdown(id='titleout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.Label('Journal'),
    html.Div([
        dcc.Input(id='journal', value='Nature', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='journalstyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='journalout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.Label('Year'),
    html.Div([
        dcc.Input(id='year', value='2020', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='yearstyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='yearout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),


    html.Label('Volume'),
    html.Div([
        dcc.Input(id='volume', value='10', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='volumestyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='volumeout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.Label('Issue'),
    html.Div([
        dcc.Input(id='issue', value='3', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='issuestyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='issueout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.Label('Page range'),
    html.Div([
        dcc.Input(id='pagerange', value='1231-1238', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='pagestyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='pageout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.Label('DOI'),
    html.Div([
        dcc.Input(id='doi', value='10.1234/a1bc23456d', type='text',
                  style={'width': '100%'}),
        dcc.Checklist(
                id='doistyles',
                options=[
                    {'label': 'Bold', 'value': 'bold'},
                    {'label': 'Italic', 'value': 'italic'}
                ],
                value=[]
            ),
            #dcc.Markdown(id='doiout')
    ],
    style={'display': 'inline-block', 'columns': '2'}),

    html.H4('Your reference template'),
    dcc.Markdown(id='example_reference'),

    dcc.Markdown(id='result_reference', style={'display': 'none'}),

    html.H4('To start upload .txt file with DOIs'),
    dcc.Upload(
        id='upload-data',
        children=html.Div([
            'Drag and Drop or ',
            html.A('Select File')
        ]),
        style={
            'width': '70%',
            'height': '60px',
            'lineHeight': '60px',
            'borderWidth': '1px',
            'borderStyle': 'dashed',
            'borderRadius': '5px',
            'textAlign': 'center',
            'margin': '10px'
        },
        # Allow multiple files to be uploaded
        multiple=False
    ),
    dcc.Markdown(id='output-data-upload'),
    html.Div(id='hidden-input-reference', style={'display': 'none'}),
],
style={'display': 'inline-block', 'width': '80%'})


def define_style(item, item_style, mode='console'):
    if 'bold' in item_style and 'italic' in item_style:
        if mode == 'console':
            return '***{}***'.format(item)
        elif mode == 'docx':
            return 'bold_italic'
    elif 'bold' in item_style:
        if mode == 'console':
            return '**{}**'.format(item)
        elif mode == 'docx':
            return 'bold'
    elif 'italic' in item_style:
        if mode == 'console':
            return '*{}*'.format(item)
        elif mode == 'docx':
            return 'italic'
    else:
        if mode == 'console':
            return '{}'.format(item)
        elif mode == 'docx':
            return 'normal'


def format_to_docx_style(item, paragraph, style):
    if 'bold' in style and 'italic' in style:
        tmp = paragraph.add_run(item)
        tmp.italic = True
        tmp.bold = True
    elif 'bold' in style:
        paragraph.add_run(item).bold = True
    elif 'italic' in style:
        paragraph.add_run(item).italic = True
    else:
        paragraph.add_run(item)


# uploadfile
def read_upload_file(filename):
    try:
        f = open(filename, 'r', encoding='utf-8')
        references_list = []
        for line in f.readlines():
            if line[-1:] == '\n':
                references_list.append(RAF.parser(line[:-1]))
            else:
                references_list.append(RAF.parser(line))
        return '### File {} upload succesfully: it contains {} references'.format(filename, len(references_list)), references_list
    except:
        return '### Something wrong with {}'.format(filename), None


# upload file
@app.callback(
    [Output(component_id='output-data-upload', component_property='children'),
     Output(component_id='hidden-input-reference', component_property='children')],
    [Input(component_id='upload-data', component_property='filename')]
)
def update_output(filename):
    if filename is not None:
        children = [read_upload_file(filename)]
        return children[0][0], children[0][1]
    else:
        return (no_update, no_update)


# example reference
@app.callback(
    Output(component_id='example_reference', component_property='children'),
    [Input(component_id='name', component_property='value'),
     Input(component_id='nametype', component_property='value'),
     Input(component_id='secname', component_property='value'),
     Input(component_id='secnametype', component_property='value'),
     Input(component_id='surname', component_property='value'),
     Input(component_id='surnametype', component_property='value'),
     Input(component_id='namesdelimiter', component_property='value'),
     Input(component_id='Title', component_property='value'),
     Input(component_id='titlestyles', component_property='value'),
     Input(component_id='journal', component_property='value'),
     Input(component_id='journalstyles', component_property='value'),
     Input(component_id='year', component_property='value'),
     Input(component_id='yearstyles', component_property='value'),
     Input(component_id='volume', component_property='value'),
     Input(component_id='volumestyles', component_property='value'),
     Input(component_id='issue', component_property='value'),
     Input(component_id='issuestyles', component_property='value'),
     Input(component_id='pagerange', component_property='value'),
     Input(component_id='pagestyles', component_property='value'),
     Input(component_id='doi', component_property='value'),
     Input(component_id='doistyles', component_property='value'),
     Input(component_id='nameorder', component_property='value'),
     Input(component_id='referenceorder', component_property='value'),
     Input(component_id='itemsdelimiter', component_property='value'),
     ]
)
def update_example_reference(name, nametype, secname, secnametype, surname, surnametype, namesdelimiter, title, titlestyles,
                          journal, journalstyles, year, yearstyles, volume, volumestyles, issue, issuestyles,
                          pagerange, pagestyles, doi, doistyles, nameorder, referenceorder, itemsdelimiter):

    ref_list = [[[name, secname, surname]], title, journal, year, volume, issue, pagerange, doi]
    namesdelimiter = namesdelimiter.split('#')
    itemsdelimiter = itemsdelimiter.split('#')

    reference = '### '
    count = 0
    for item in referenceorder:
        countname = 0
        if item == 'NAMES':
            for author in ref_list[0]:
                for flag in nameorder:
                    if flag == 'NAME':
                        if nametype == 'full':
                            reference += author[0]
                        elif nametype == 'short':
                            reference += author[0][:1] + '.'
                    if flag == 'SECNAME':
                        if secnametype == 'full':
                            reference += author[1]
                        elif secnametype == 'short':
                            reference += author[1][:1] + '.'
                    if flag == 'SURNAME':
                        if surnametype == 'full':
                            reference += author[2]
                        elif surnametype == 'short':
                            reference += author[2][:1] + '.'
                    try:
                        reference += namesdelimiter[countname]
                    except:
                        reference += ' '
                    countname += 1

        if item == 'TITLE':
            reference += define_style(ref_list[1], titlestyles)
        if item == 'JOURNAL':
            reference += define_style(ref_list[2], journalstyles)
        if item == 'YEAR':
            reference += define_style(ref_list[3], yearstyles)
        if item == 'VOLUME':
            reference += define_style(ref_list[4], volumestyles)
        if item == 'ISSUE':
            reference += define_style(ref_list[5], issuestyles)
        if item == 'PAGERANGE':
            reference += define_style(ref_list[6], pagestyles)
        if item == 'DOI':
            reference += define_style(ref_list[7], doistyles)
        try:
            reference += itemsdelimiter[count]
        except:
            reference += ' '
        count += 1

    return reference


# parsed references
@app.callback(
    Output(component_id='result_reference', component_property='children'),
    [Input(component_id='hidden-input-reference', component_property='children'), #####################
     Input(component_id='nametype', component_property='value'),
     Input(component_id='secnametype', component_property='value'),
     Input(component_id='surnametype', component_property='value'),
     Input(component_id='namesdelimiter', component_property='value'),
     Input(component_id='titlestyles', component_property='value'),
     Input(component_id='journalstyles', component_property='value'),
     Input(component_id='yearstyles', component_property='value'),
     Input(component_id='volumestyles', component_property='value'),
     Input(component_id='issuestyles', component_property='value'),
     Input(component_id='pagestyles', component_property='value'),
     Input(component_id='doistyles', component_property='value'),
     Input(component_id='nameorder', component_property='value'),
     Input(component_id='referenceorder', component_property='value'),
     Input(component_id='itemsdelimiter', component_property='value'),
     ]
)
def update_parser_reference(URL_list, nametype, secnametype, surnametype, namesdelimiter, titlestyles,
                          journalstyles, yearstyles, volumestyles, issuestyles, pagestyles, doistyles, nameorder,
                          referenceorder, itemsdelimiter):

    def form_reference(input_reference, nametype, secnametype, surnametype, namesdelimiter, titlestyles,
                          journalstyles, yearstyles, volumestyles, issuestyles, pagestyles, doistyles, nameorder,
                          referenceorder, itemsdelimiter):
        namesdelimiter = namesdelimiter.split('#')
        itemsdelimiter = itemsdelimiter.split('#')
        name_delimiters = dict(zip(nameorder, namesdelimiter))
        count = 0
        #print(nameorder)
        for item in referenceorder:
            if item == 'NAMES':
                for author in input_reference[0]:
                    countname = 0
                    if author is not input_reference[0][-1]:  # если автор не последний в статье
                        noskip = []
                        tmpnameorder = []
                        namesdelimiter_short = []
                        if author[0] != '-':
                            noskip.append('NAME')
                        if author[1] != '-':
                            noskip.append('SECNAME')
                        if author[2] != '-':
                            noskip.append('SURNAME')

                        for item in nameorder:
                            if item in noskip:
                                tmpnameorder.append(item)
                                namesdelimiter_short.append(name_delimiters[item])

                        if len(tmpnameorder) == 2:
                            namesdelimiter_short.pop()
                            namesdelimiter_short.append(name_delimiters[nameorder[-1]])

                        for item in range(len(tmpnameorder)):
                            if tmpnameorder[item] == 'NAME':
                                if nametype == 'full':
                                    par1.add_run(author[0])
                                elif nametype == 'short':
                                    par1.add_run(author[0][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])

                            if tmpnameorder[item] == 'SECNAME':
                                if secnametype == 'full':
                                    par1.add_run(author[1])
                                elif secnametype == 'short':
                                    par1.add_run(author[1][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])

                            if tmpnameorder[item] == 'SURNAME':
                                if surnametype == 'full':
                                    par1.add_run(author[2])
                                elif surnametype == 'short':
                                    par1.add_run(author[2][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])
                    else:
                        noskip = []
                        tmpnameorder = []
                        namesdelimiter_short = []
                        if author[0] != '-':
                            noskip.append('NAME')
                        if author[1] != '-':
                            noskip.append('SECNAME')
                        if author[2] != '-':
                            noskip.append('SURNAME')

                        for item in nameorder:
                            if item in noskip:
                                tmpnameorder.append(item)
                                namesdelimiter_short.append(name_delimiters[item])

                        namesdelimiter_short.pop()
                        namesdelimiter_short.append('')

                        for item in range(len(tmpnameorder)):
                            if tmpnameorder[item] == 'NAME':
                                if nametype == 'full':
                                    par1.add_run(author[0])
                                elif nametype == 'short':
                                    par1.add_run(author[0][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])

                            if tmpnameorder[item] == 'SECNAME':
                                if secnametype == 'full':
                                    par1.add_run(author[1])
                                elif secnametype == 'short':
                                    par1.add_run(author[1][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])

                            if tmpnameorder[item] == 'SURNAME':
                                if surnametype == 'full':
                                    par1.add_run(author[2])
                                elif surnametype == 'short':
                                    par1.add_run(author[2][:1] + '.')
                                par1.add_run(namesdelimiter_short[item])

            if item == 'TITLE':
                format_to_docx_style(input_reference[1], par1, titlestyles)
            if item == 'JOURNAL':
                format_to_docx_style(input_reference[2], par1, journalstyles)
            if item == 'YEAR':
                format_to_docx_style(input_reference[3], par1, yearstyles)
            if item == 'VOLUME':
                format_to_docx_style(input_reference[4], par1, volumestyles)
            if item == 'ISSUE':
                format_to_docx_style(input_reference[5], par1, issuestyles)
            if item == 'PAGERANGE':
                format_to_docx_style(input_reference[6], par1, pagestyles)
            if item == 'DOI':
                format_to_docx_style(input_reference[7], par1, doistyles)
            try:
                par1.add_run(itemsdelimiter[count])
            except:
                par1.add_run(' ')
            count += 1

    try:
        doc = docx.Document()
        par1 = doc.add_paragraph('References:\n\n')
        for input_reference in URL_list:
            form_reference(input_reference, nametype, secnametype, surnametype, namesdelimiter, titlestyles,
                          journalstyles, yearstyles, volumestyles, issuestyles, pagestyles, doistyles, nameorder,
                          referenceorder, itemsdelimiter)
            par1.add_run('\n\n')

        doc.save('ReFFormator_refereces.docx')
        return 'Success'
    except:
        URL = ''


if __name__ == '__main__':
    app.run_server(debug=True)
